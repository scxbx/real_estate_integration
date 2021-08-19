import os

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from docx.shared import RGBColor
from tkinter import filedialog


def mod_word(in_path, out_path, num):
    document = Document(in_path)
    tables = document.tables

    n_cell = tables[0].cell(2, 1)
    tl = n_cell.text.split('-')
    n_cell.text = tl[0] + '-' + str("%04d" % num)
    print(tables[0].cell(2, 1).text)
    cell_par = n_cell.paragraphs[0]
    cell_par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 获取 run 对象
    cell_run = n_cell.paragraphs[0].runs[0]
    # 设置字体
    cell_run.font.name = '仿宋_GB2312'
    cell_run.font.size = Pt(14)

    document.save(out_path)


if __name__ == '__main__':
    # mod_word(r'C:\Users\sc\Documents\飞行者科技\房地一体\曾广深\指界通知书光坡镇.docx')
    #
    # for i in range(1, 10000):
    #     mod_word()

    print("%04d" % 1)
    inpath = filedialog.askopenfilename()
    outpath = os.path.join(os.getcwd(), 'results')

    for i in range(1, 10000):
        mod_word(inpath, os.path.join(outpath, str(i) + '.docx'), i)


